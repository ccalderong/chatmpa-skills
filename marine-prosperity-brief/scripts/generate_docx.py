#!/usr/bin/env python3
"""Convert policy brief markdown files to formatted DOCX with embedded maps."""

import os
import re
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from chatmpa.brand import COLORS, FONTS, LOGO_PATH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_DIR = BASE_DIR
MAPS_DIR = os.path.join(BASE_DIR, 'maps')
OUT_DIR = os.path.join(BASE_DIR, 'documents_docx')
FIGURES_DIR = os.path.join(os.path.dirname(BASE_DIR), 'figures')
os.makedirs(OUT_DIR, exist_ok=True)

PROSPERITY_FIG_PATH = os.path.join(BASE_DIR, 'figures', 'prosperity_3panel.png')
PROSPERITY_FIG_CAPTION = (
    'Figure 1 | Marine Prosperity Index: National Distribution Across Mexico\'s Coast. '
    '(a) Balance — evenness-based coordination metric (B ∈ [0,1]), where B = 1 indicates '
    'perfect coordination across Nature, Livelihood, and Well-being axes. High-balance cells '
    '(mako scale, dark = low, light = high) are widely distributed along the coast. '
    '(b) Level — overall performance metric L = (Nature + Livelihood + Well-being) / 3, '
    'capturing whether communities are prospering at a high absolute standard (cividis scale). '
    '(c) Prosperity Pp = Balance × Level — the composite indicator integrating both coordination '
    'and performance (inferno scale). National means: Balance = 0.80, Level = 0.37, Prosperity = 0.30. '
    'Livelihood is the binding constraint in 88% of coastal cells, making it the dominant driver '
    'of low Prosperity scores across Mexico\'s coast.'
)

FRAMEWORK_SUMMARY = [
    ('The Marine Prosperity Index (MPpI) Framework',),   # title only
    ('The MPpI is a diagnostic tool that classifies coastal communities by both the '
     'balance across three development dimensions and the overall level of performance. '
     'Unlike frameworks that treat dimensions independently, the MPpI identifies which '
     'axis constrains a system and sequences investment accordingly.', 'body'),
    ('Three Dimensions', 'h3'),
    ('**Nature** (13 indicators): Biodiversity, ecosystem productivity, water quality, '
     'carbon storage, and climate stress.', 'bullet'),
    ('**Livelihood** (12 indicators): Income, employment, investment, and fisheries '
     'performance.', 'bullet'),
    ('**Well-being** (23 indicators): Education, health, household services, poverty, '
     'and governance.', 'bullet'),
    ('Key Metrics', 'h3'),
    ('**Balance (B) — evenness-based metric, B ∈ [0, 1].** Computed as '
     'B = (E − 1/3)/(2/3), where E = (Σxᵢ)²/(n·Σxᵢ²). '
     'High balance (B ≥ 0.75) means no single dimension severely lags behind.', 'bullet'),
    ('**Level (L) = (Nature + Livelihood + Well-being) / 3.** Measures overall performance. '
     'High level (L ≥ 0.40) indicates strong development across the board.', 'bullet'),
    ('**Prosperity (Pp = Balance × Level).** Composite indicator integrating coordination '
     'and overall performance into a single score. National coastal mean Pp ≈ 0.30.', 'bullet'),
    ('**Viability threshold:** Any axis score below 0.20 requires priority rescue intervention, '
     'regardless of balance status.', 'bullet'),
    ('Four Prosperity Categories', 'h3'),
    ('**Balanced Prosperity** (high balance, high level): Coordinated development — maintain '
     'trajectory and strengthen the limiting axis.', 'bullet'),
    ('**Imbalanced Prosperity** (low balance, high level): Strong overall but one axis lags — '
     'target investment in the binding constraint.', 'bullet'),
    ('**Balanced Underdevelopment** (high balance, low level): Uniformly constrained — broad '
     'investment across all axes needed.', 'bullet'),
    ('**Imbalanced Underdevelopment** (low balance, low level): Severe deficits with one axis '
     'particularly weak — urgent priority intervention required.', 'bullet'),
    ('Binding Constraint Logic', 'h3'),
    ('Nationally, livelihood is the binding constraint in 88% of coastal grid cells, nature in '
     '9% (industrialized ports, high-tourism zones), and well-being in 3% (remote indigenous '
     'communities). Investment in a dimension that is already adequate cannot improve balance; '
     'only targeting the weakest axis closes the prosperity gap efficiently.', 'body'),
]

def _hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _hex_to_RGBColor(hex_color):
    return RGBColor(*_hex_to_rgb(hex_color))


COLOR_PRIMARY   = _hex_to_RGBColor(COLORS["deep_sea"])
COLOR_ACCENT    = _hex_to_RGBColor(COLORS["ocean_blue"])
COLOR_TEXT_DARK = _hex_to_RGBColor(COLORS["text_dark"])
COLOR_TEXT_GRAY = _hex_to_RGBColor(COLORS["text_gray"])
COLOR_HEADER_BG = _hex_to_rgb(COLORS["deep_sea"])
COLOR_ALT_BG    = _hex_to_rgb(COLORS["light_blue"])
_DEEP_SEA_HEX   = COLORS["deep_sea"].lstrip('#')
_OCEAN_BLUE_HEX = COLORS["ocean_blue"].lstrip('#')
_LIGHT_BLUE_HEX = COLORS["light_blue"].lstrip('#')

COMMUNITY_MAP = {
    'alto_golfo':              ('policy_brief_alto_golfo.md',              'alto_golfo'),
    'bahia_de_banderas':       ('policy_brief_bahia_de_banderas.md',       'bahia_de_banderas'),
    'bahia_de_los_angeles':    ('policy_brief_bahia_de_los_angeles.md',    'bahia_de_los_angeles'),
    'bahia_de_kino':           ('policy_brief_bahia_de_kino.md',            'bahia_de_kino'),
    'el_manglito':             ('policy_brief_el_manglito.md',             'el_manglito'),
    'la_manga':                ('policy_brief_la_manga.md',                'la_manga'),
    'la_reforma':              ('policy_brief_la_reforma.md',              'la_reforma'),
    'la_ribera':               ('policy_brief_la_ribera.md',               'la_ribera'),
    'punta_chueca':            ('policy_brief_punta_chueca.md',            'punta_chueca'),
    'san_basilio':             ('policy_brief_san_basilio.md',             'san_basilio'),
    'san_carlos':              ('policy_brief_san_carlos.md',              'san_carlos'),
}


# ─── Document helpers ────────────────────────────────────────────────────────

def set_page_margins(doc, top=1, bottom=1, left=1, right=1):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def set_cell_background(cell, hex_rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '%02X%02X%02X' % tuple(hex_rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_TEXT_GRAY


def set_paragraph_spacing(para, space_before=0, space_after=6, line_spacing=1.15):
    from docx.shared import Pt
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    # line spacing (240 = single, 276 = 1.15)
    spacing.set(qn('w:line'), str(int(line_spacing * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


# ─── Content parsers ─────────────────────────────────────────────────────────

def parse_inline(text):
    """Return list of (run_text, bold, italic) tuples from markdown inline markup."""
    segments = []
    pattern = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False, False))
        if m.group(1):
            segments.append((m.group(1), True, True))
        elif m.group(2):
            segments.append((m.group(2), True, False))
        elif m.group(3):
            segments.append((m.group(3), False, True))
        elif m.group(4):
            segments.append((m.group(4), False, False))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False, False))
    return segments


def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                             font_size=11, color=None, align=None,
                             space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, space_before, space_after)
    if align:
        p.alignment = align

    for seg_text, seg_bold, seg_italic in parse_inline(text):
        run = p.add_run(seg_text)
        run.bold = bold or seg_bold
        run.italic = italic or seg_italic
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
    return p


def parse_table_markdown(lines):
    """Parse markdown table lines → (headers, rows)."""
    headers = []
    rows = []
    for line in lines:
        line = line.strip()
        if not line or re.match(r'^\|[-| :]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    # Header row
    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_background(cell, COLOR_HEADER_BG)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = COLOR_ALT_BG if ri % 2 == 1 else (0xFF, 0xFF, 0xFF)
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            for seg, bold, italic in parse_inline(cell_text):
                run = p.add_run(seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(9.5)

    return table


# ─── DOCX builder ────────────────────────────────────────────────────────────

def build_docx(md_path, map_slug, community_name):
    with open(md_path, encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()
    set_page_margins(doc)
    add_page_numbers(doc)

    # ── Styles ──
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = FONTS["body"]
    normal.font.size = Pt(11)

    # ── Header block ──
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(logo_para, 0, 8)
        logo_para.add_run().add_picture(LOGO_PATH, width=Inches(2.0))

    # Pull title from first H1
    title_text = community_name
    for line in lines:
        if line.startswith('# '):
            title_text = line[2:].strip()
            break

    # Title
    p_title = doc.add_paragraph()
    set_paragraph_spacing(p_title, 0, 4)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.name = FONTS["display"]

    # Horizontal rule (thin blue paragraph border)
    p_rule = doc.add_paragraph()
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), _DEEP_SEA_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_paragraph_spacing(p_rule, 0, 8)

    # Meta info from **Key:** Value lines
    for line in lines[1:20]:
        line = line.strip()
        if line.startswith('**') and ':' in line and not line.startswith('##'):
            p_meta = doc.add_paragraph()
            set_paragraph_spacing(p_meta, 0, 3)
            for seg, bold, italic in parse_inline(line):
                run = p_meta.add_run(seg)
                run.bold = bold
                run.font.size = Pt(10.5)
                if bold:
                    run.font.color.rgb = COLOR_PRIMARY

    # Separator
    doc.add_paragraph()

    # ── MPpI Framework Summary ──
    _add_section_h2(doc, 'About the Marine Prosperity Index')
    for item in FRAMEWORK_SUMMARY[1:]:  # skip title-only tuple
        text = item[0]
        kind = item[1] if len(item) > 1 else 'body'
        if kind == 'h3':
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 8, 2)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_ACCENT
        elif kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            set_paragraph_spacing(p, 0, 2)
            for seg, bold, italic in parse_inline(text):
                run = p.add_run(seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 0, 5)
            for seg, bold, italic in parse_inline(text):
                run = p.add_run(seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10.5)

    # ── Figure 1 (framework application) ──
    if os.path.exists(PROSPERITY_FIG_PATH):
        p_fig3 = doc.add_paragraph()
        p_fig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fig3.add_run().add_picture(PROSPERITY_FIG_PATH, width=Inches(6.3))
        set_paragraph_spacing(p_fig3, 0, 4)

        p_cap = doc.add_paragraph()
        set_paragraph_spacing(p_cap, 0, 12)
        run_cap = p_cap.add_run(PROSPERITY_FIG_CAPTION)
        run_cap.italic = True
        run_cap.font.size = Pt(9)
        run_cap.font.color.rgb = COLOR_TEXT_DARK

    # ── Location map ──
    map_path = os.path.join(MAPS_DIR, f'{map_slug}_location_map.png')
    if os.path.exists(map_path):
        _add_section_h2(doc, 'Community Location')

        p_map = doc.add_paragraph()
        p_map.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_map.add_run()
        run_img.add_picture(map_path, width=Inches(6.3))
        set_paragraph_spacing(p_map, 0, 12)

    # ── Parse and render body ──
    i = 0
    in_table = False
    table_lines = []

    # Skip header lines already processed (until first ##)
    start_idx = 0
    for idx, line in enumerate(lines):
        if line.startswith('## '):
            start_idx = idx
            break

    i = start_idx
    while i < len(lines):
        line = lines[i]

        # H2 section header
        if line.startswith('## '):
            if in_table:
                headers, rows = parse_table_markdown(table_lines)
                add_table(doc, headers, rows)
                doc.add_paragraph()
                in_table = False
                table_lines = []

            text = line[3:].strip()
            _add_section_h2(doc, text)

        # H3 sub-header
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 8, 2)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_ACCENT

        # H4 sub-sub-header
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 5, 2)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_TEXT_DARK

        # Horizontal rule
        elif line.strip() == '---':
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 3, 3)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bdr = OxmlElement('w:bottom')
            bdr.set(qn('w:val'), 'single')
            bdr.set(qn('w:sz'), '4')
            bdr.set(qn('w:space'), '1')
            bdr.set(qn('w:color'), _LIGHT_BLUE_HEX)
            pBdr.append(bdr)
            pPr.append(pBdr)

        # Table line
        elif line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)

        # End of table
        elif in_table and not line.strip().startswith('|'):
            headers, rows = parse_table_markdown(table_lines)
            add_table(doc, headers, rows)
            p_gap = doc.add_paragraph()
            set_paragraph_spacing(p_gap, 0, 6)
            in_table = False
            table_lines = []

            # Process current line normally (fall through)
            if line.strip():
                _add_body_line(doc, line)

        # Bullet / numbered list
        elif line.strip().startswith('- ') or re.match(r'^\d+\. ', line.strip()):
            bullet_text = _fix_grid_text(re.sub(r'^[-\d]+[.)]\s*', '', line.strip()))
            p = doc.add_paragraph(style='List Bullet')
            set_paragraph_spacing(p, 0, 2)
            for seg, bold, italic in parse_inline(bullet_text):
                run = p.add_run(seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10.5)

        # Sub-bullet
        elif line.strip().startswith('  - '):
            bullet_text = _fix_grid_text(line.strip()[4:])
            p = doc.add_paragraph(style='List Bullet 2')
            set_paragraph_spacing(p, 0, 2)
            for seg, bold, italic in parse_inline(bullet_text):
                run = p.add_run(seg)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(10.5)

        # Empty line
        elif line.strip() == '':
            pass  # skip blank lines (paragraph spacing handles gaps)

        # Body text
        else:
            _add_body_line(doc, line)

        i += 1

    # Flush any trailing table
    if in_table and table_lines:
        headers, rows = parse_table_markdown(table_lines)
        add_table(doc, headers, rows)

    return doc


def _add_section_h2(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 14, 4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.name = FONTS["display"]
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement('w:bottom')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '4')
    bdr.set(qn('w:space'), '1')
    bdr.set(qn('w:color'), _OCEAN_BLUE_HEX)
    pBdr.append(bdr)
    pPr.append(pBdr)
    return p


def _fix_grid_text(text):
    """Correct grid cell resolution references in body text."""
    text = re.sub(r'0\.05°\s*\(~?5\s*km\)', '0.1° × 0.1° (~11 km, ~121 km²)', text)
    text = re.sub(r'0\.05°\s*resolution\s*\(~?5\.?5?\s*km\)', '0.1° × 0.1° resolution (~121 km²)', text)
    text = re.sub(r'~5\.5\s*km\)', '~121 km²)', text)
    return text


def _add_body_line(doc, line):
    line = _fix_grid_text(line.strip())
    if not line:
        return
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 5)
    for seg, bold, italic in parse_inline(line):
        run = p.add_run(seg)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10.5)


# ─── Consolidated report ─────────────────────────────────────────────────────

COMMUNITIES_DATA = {
    'Alto Golfo':           {'nature': 0.51, 'livelihood': 0.42, 'wellbeing': 0.49, 'balance': 0.87, 'level': 0.47, 'prosperity': 0.41, 'category': 'Balanced Prosperity',     'limiting': 'Livelihood', 'cells': 27,  'state': 'Baja California'},
    'Bahía de Banderas':    {'nature': 0.43, 'livelihood': 0.18, 'wellbeing': 0.46, 'balance': 0.80, 'level': 0.36, 'prosperity': 0.29, 'category': 'Balanced but Developing', 'limiting': 'Livelihood', 'cells': 20,  'state': 'Nayarit / Jalisco'},
    'Bahía de Los Ángeles': {'nature': 0.48, 'livelihood': 0.75, 'wellbeing': 0.52, 'balance': 0.94, 'level': 0.58, 'prosperity': 0.55, 'category': 'Balanced Prosperity',     'limiting': 'Nature',     'cells': 17,  'state': 'Baja California'},
    'Bahía de Kino':        {'nature': 0.46, 'livelihood': 0.84, 'wellbeing': 0.60, 'balance': 0.91, 'level': 0.63, 'prosperity': 0.58, 'category': 'Balanced Prosperity',     'limiting': 'Nature',     'cells': 23,  'state': 'Sonora'},
    'El Manglito':          {'nature': 0.40, 'livelihood': 0.25, 'wellbeing': 0.48, 'balance': 0.90, 'level': 0.37, 'prosperity': 0.34, 'category': 'Balanced but Developing', 'limiting': 'Livelihood', 'cells': 20,  'state': 'Baja California Sur'},
    'La Manga':             {'nature': 0.54, 'livelihood': 0.35, 'wellbeing': 0.47, 'balance': 0.95, 'level': 0.45, 'prosperity': 0.43, 'category': 'Balanced Prosperity',     'limiting': 'Livelihood', 'cells': 24,  'state': 'Sonora'},
    'La Reforma':           {'nature': 0.44, 'livelihood': 0.15, 'wellbeing': 0.52, 'balance': 0.74, 'level': 0.37, 'prosperity': 0.27, 'category': 'Lagging',                 'limiting': 'Livelihood', 'cells': 16,  'state': 'Sinaloa'},
    'La Ribera':            {'nature': 0.38, 'livelihood': 0.34, 'wellbeing': 0.49, 'balance': 0.96, 'level': 0.40, 'prosperity': 0.38, 'category': 'Balanced but Developing', 'limiting': 'Livelihood', 'cells': 21,  'state': 'Baja California Sur'},
    'Punta Chueca':         {'nature': 0.46, 'livelihood': 0.80, 'wellbeing': 0.59, 'balance': 0.91, 'level': 0.62, 'prosperity': 0.56, 'category': 'Balanced Prosperity',     'limiting': 'Nature',     'cells': 20,  'state': 'Sonora'},
    'San Basilio':          {'nature': 0.38, 'livelihood': 0.11, 'wellbeing': 0.49, 'balance': 0.69, 'level': 0.33, 'prosperity': 0.22, 'category': 'Lagging',                 'limiting': 'Livelihood', 'cells': 23,  'state': 'Sonora'},
    'San Carlos':           {'nature': 0.54, 'livelihood': 0.35, 'wellbeing': 0.46, 'balance': 0.95, 'level': 0.45, 'prosperity': 0.43, 'category': 'Balanced Prosperity',     'limiting': 'Livelihood', 'cells': 24,  'state': 'Sonora'},
}


def build_consolidated(output_path):
    doc = Document()
    set_page_margins(doc)
    add_page_numbers(doc)

    # ── Title page ──
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(logo_para, 36, 12)
        logo_para.add_run().add_picture(LOGO_PATH, width=Inches(2.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 6)
    run = p.add_run('Marine Prosperity Index')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, 0, 6)
    run2 = p2.add_run('Gulf of California Regional Assessment')
    run2.bold = True
    run2.font.size = Pt(18)
    run2.font.color.rgb = COLOR_ACCENT

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p3, 12, 4)
    run3 = p3.add_run('May 2026 | Marine Prosperity Index (MPpI) Framework')
    run3.font.size = Pt(11)
    run3.font.color.rgb = COLOR_TEXT_GRAY

    doc.add_page_break()

    # ── Section 1: Overview ──
    def h2(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 14, 4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_PRIMARY
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bdr = OxmlElement('w:bottom')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '4')
        bdr.set(qn('w:space'), '1')
        bdr.set(qn('w:color'), _OCEAN_BLUE_HEX)
        pBdr.append(bdr)
        pPr.append(pBdr)

    def h3(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 8, 2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_ACCENT

    def body(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 5)
        for seg, bold, italic in parse_inline(text):
            run = p.add_run(seg)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(10.5)

    def bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        set_paragraph_spacing(p, 0, 2)
        for seg, bold, italic in parse_inline(text):
            run = p.add_run(seg)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(10.5)

    h2('1. Introduction and Regional Context')
    body('The Gulf of California (Sea of Cortez) is one of the world\'s most biologically '
         'productive and biodiverse marine regions, supporting critical fisheries, endemic '
         'wildlife, and coastal communities that depend on healthy marine ecosystems. This '
         'regional assessment applies the Marine Prosperity Index (MPpI) framework to ten '
         'coastal communities spanning the Gulf\'s eastern and western shores — from the '
         'hyper-arid upper gulf to the subtropical southern tip of the Baja California Peninsula.')
    body('The MPpI evaluates coastal systems across three dimensions: **Nature** (ecological '
         'integrity and ecosystem services), **Livelihood** (economic opportunity and fisheries '
         'performance), and **Well-being** (social services, health, education, and governance). '
         'Together, these axes capture whether a community is truly prospering in a balanced, '
         'sustainable way — or whether growth in one dimension masks critical deficits in others.')

    h3('Why the Gulf of California?')
    bullet('**Ecological significance:** Recognized by UNESCO as a World Heritage Site; home to '
           '39 cetacean species, 891 fish species, and the critically endangered vaquita porpoise.')
    bullet('**Economic importance:** Supports approximately 30% of Mexico\'s total fisheries '
           'production and anchors coastal livelihoods for hundreds of thousands of people.')
    bullet('**Regional diversity:** The ten communities span four states (Baja California, '
           'Baja California Sur, Sonora, Sinaloa) and represent widely varying prosperity profiles '
           '— from thriving balanced communities to severely constrained underdeveloped areas.')
    bullet('**Policy urgency:** Climate change, overfishing, and urban expansion are accelerating '
           'pressures on gulf ecosystems, making evidence-based investment prioritization critical.')

    h2('2. Framework Summary')
    body('The MPpI classifies each coastal grid cell (0.1° × 0.1°, ~121 km²) into one of four '
         'prosperity categories based on two key metrics:')
    bullet('**Balance (B):** Evenness-based coordination metric, '
           'B = (E − 1/3)/(2/3) where E = (Σxᵢ)²/(n·Σxᵢ²), rescaled to [0, 1]. '
           'High balance (B ≥ 0.75) indicates no single dimension is severely lagging.')
    bullet('**Level (L):** Overall performance, L = (Nature + Livelihood + Well-being) / 3. '
           'High level (L ≥ 0.40) indicates strong overall development.')
    bullet('**Prosperity (Pp = Balance × Level):** Composite indicator integrating coordination '
           'and performance. National coastal mean Pp ≈ 0.30.')

    headers = ['Category', 'Balance', 'Level', 'Description', 'Policy Approach']
    rows = [
        ['Balanced Prosperity', '≥ 0.70', '≥ 0.40', 'Coordinated, high-performing', 'Maintain + optimize limiting axis'],
        ['Imbalanced Prosperity', '< 0.70', '≥ 0.40', 'Strong but one axis lags', 'Targeted investment in lagging axis'],
        ['Balanced Underdevelopment', '≥ 0.70', '< 0.40', 'Uniform low performance', 'Broad investment across all axes'],
        ['Imbalanced Underdevelopment', '< 0.70', '< 0.40', 'Severe deficits, one worst', 'Prioritize binding constraint urgently'],
    ]
    add_table(doc, headers, rows, col_widths=[1.6, 0.7, 0.7, 1.8, 2.0])
    doc.add_paragraph()

    # ── Figure 1: inserted after framework explanation ──
    if os.path.exists(PROSPERITY_FIG_PATH):
        p_fig3 = doc.add_paragraph()
        p_fig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fig3.add_run().add_picture(PROSPERITY_FIG_PATH, width=Inches(6.3))
        set_paragraph_spacing(p_fig3, 0, 4)

        p_cap = doc.add_paragraph()
        set_paragraph_spacing(p_cap, 0, 14)
        run_cap = p_cap.add_run(PROSPERITY_FIG_CAPTION)
        run_cap.italic = True
        run_cap.font.size = Pt(9)
        run_cap.font.color.rgb = COLOR_TEXT_DARK

    h2('3. Regional Summary: All Ten Communities')
    body('The table below presents the MPpI scores for all ten Gulf of California communities '
         'assessed in this report, alongside national coastal averages for comparison.')

    nat_avg = 0.44
    liv_avg = 0.23
    wb_avg = 0.44
    bal_avg = 0.80
    lev_avg = 0.37
    pp_avg  = 0.30

    headers2 = ['Community', 'State', 'Nature', 'Livelihood', 'Well-being',
                'Balance', 'Level', 'Prosperity', 'Category', 'Limiting Axis']
    rows2 = []
    for name, d in COMMUNITIES_DATA.items():
        rows2.append([
            name, d['state'],
            f"{d['nature']:.2f}", f"{d['livelihood']:.2f}", f"{d['wellbeing']:.2f}",
            f"{d['balance']:.2f}", f"{d['level']:.2f}", f"{d['prosperity']:.2f}",
            d['category'], d['limiting']
        ])
    rows2.append(['**National Average**', '—',
                  f'{nat_avg:.2f}', f'{liv_avg:.2f}', f'{wb_avg:.2f}',
                  f'{bal_avg:.2f}', f'{lev_avg:.2f}', f'{pp_avg:.2f}', '—', '—'])
    add_table(doc, headers2, rows2,
              col_widths=[1.3, 0.85, 0.50, 0.70, 0.70, 0.55, 0.50, 0.65, 1.20, 0.75])
    doc.add_paragraph()

    h2('4. Regional Patterns and Analysis')

    h3('4.1 Prosperity Category Distribution')
    cats = {}
    for d in COMMUNITIES_DATA.values():
        cats[d['category']] = cats.get(d['category'], 0) + 1
    for cat, n in sorted(cats.items(), key=lambda x: -x[1]):
        bullet(f'**{cat}:** {n} communities ({int(n/10*100)}%)')

    body('The Gulf shows a predominantly **livelihood-constrained** profile, consistent with '
         'national patterns (88% of coastal cells are livelihood-limited). However, the region '
         'also includes communities where **Nature** is the binding constraint — a pattern '
         'associated with high-tourism or industrialized economies where economic development '
         'outpaces ecological conditions.')

    h3('4.2 North–South Gradient')
    body('A clear north-to-south gradient is visible in the data:')
    bullet('**Upper Gulf (San Felipe, Bahía de Los Ángeles):** Ecologically rich but economically '
           'challenged; the critically endangered vaquita porpoise makes this zone a global '
           'conservation priority.')
    bullet('**Mid-Gulf (Punta Chueca, San Carlos, La Manga, San Basilio):** Sonoran communities '
           'show high economic activity in some locations (Punta Chueca: Livelihood = 0.80) but '
           'significant variation in ecological and social conditions.')
    bullet('**Lower Gulf / BCS (El Manglito, La Ribera, San Quintín):** Communities near La Paz '
           'and Los Cabos show strong social and environmental scores but face tourism-driven '
           'livelihood transitions and coastal development pressures.')
    bullet('**Pacific Coast / Sinaloa (La Reforma):** Severely livelihood-constrained despite '
           'adequate nature and well-being scores; intensive fishing pressure and limited '
           'economic diversification dominate the regional picture.')

    h3('4.3 Nature-Limited vs. Livelihood-Limited Communities')
    body('Two communities — **Bahía de Los Ángeles** and **Punta Chueca** — stand out as '
         'nature-limited, meaning their economic prosperity has outpaced ecological conditions. '
         'These require conservation investment as the primary policy response.')
    body('The remaining eight communities are livelihood-limited, calling for targeted economic '
         'interventions: fisheries value chains, aquaculture, ecotourism, and employment programs.')

    h2('5. Regional Policy Recommendations')

    h3('5.1 Cross-Community Priorities')
    bullet('**Fisheries value chain investment:** Across all ten communities, cold storage, '
           'processing facilities, and direct market access for small-scale fishers represent '
           'the highest-leverage livelihood intervention.')
    bullet('**Conservation coordination:** Nature-limited communities (Bahía de Los Ángeles, '
           'Punta Chueca) require coordinated MPA management and ecological monitoring.')
    bullet('**Tourism governance:** Several communities (La Ribera, San Carlos, San Quintín) '
           'face rapid tourism development — governance frameworks to ensure local benefit-sharing '
           'are essential.')
    bullet('**Indigenous rights:** Communities including San Felipe (Cucapá), Punta Chueca '
           '(Seri/Comcáac), and Bahía de Los Ángeles (Cochimí) host indigenous peoples whose '
           'fishing rights and territorial sovereignty must be central to any intervention design.')

    h3('5.2 Investment Sequencing')
    headers3 = ['Priority Tier', 'Communities', 'Primary Action', 'Rationale']
    rows3 = [
        ['Critical (Lagging)', 'San Basilio, La Reforma', 'Emergency livelihood investment', 'Livelihood < 0.20; below viability threshold'],
        ['High (Lagging)', 'El Manglito', 'Targeted livelihood programs', 'Livelihood-limited; L at national average'],
        ['Medium (Imbalanced)', 'San Felipe, Punta Chueca, Bahía de los Ángeles', 'Axis-specific investment', 'Imbalanced growth; one axis lagging'],
        ['Maintain (Balanced)', 'La Manga, San Carlos, San Quintín, La Ribera', 'Livelihood strengthening', 'Balanced but livelihood limiting'],
    ]
    add_table(doc, headers3, rows3, col_widths=[1.3, 2.0, 1.8, 2.1])
    doc.add_paragraph()

    h2('6. Equity and Governance Considerations')
    body('The MPpI identifies **what** to invest in across dimensions but does not prescribe '
         'how to ensure benefits reach marginalized groups. Regional equity priorities include:')
    bullet('**Indigenous coastal communities:** Cucapá (San Felipe), Seri/Comcáac (Punta Chueca), '
           'and Cochimí (Bahía de Los Ángeles) hold traditional fishing rights requiring formal '
           'recognition and protection in all development interventions.')
    bullet('**Small-scale fishers:** All ten communities have artisanal fishing sectors that are '
           'often underserved by infrastructure investments that primarily benefit established '
           'cooperatives and commercial operators.')
    bullet('**Tourism displacement risk:** Communities experiencing rapid tourism growth (La Ribera, '
           'San Carlos) face housing price increases, fishing access restrictions, and labor '
           'informality that can reverse well-being gains.')
    bullet('**Gender and youth:** Economic programs must explicitly target women and youth, who '
           'bear disproportionate livelihood constraints in fisheries-dependent communities.')
    body('**Recommendation:** Apply the Ocean Equity Index (OEI) assessment prior to implementing '
         'any MPpI-guided interventions in this region.')

    h2('7. Data Sources and Methodology')
    body('Community scores are derived from the Marine Prosperity Index dataset covering '
         '3,236 coastal grid cells across Mexico\'s coast. For each community, scores represent '
         'the average across all grid cells within a 30 km buffer of the community centroid.')
    bullet('Environmental indicators: MODIS satellite imagery, Copernicus Marine Service, CONABIO')
    bullet('Economic indicators: INEGI Economic Census (2019), CONAPESCA fisheries statistics')
    bullet('Social indicators: INEGI Population Census 2020, CONEVAL poverty measures')
    bullet('Spatial resolution: 0.05° grid cells (~25 km²)')
    bullet('Reference: Favoretto et al. (in preparation). *The Marine Prosperity Index: A Decision '
           'Framework for Balanced Coastal Development.*')

    doc.add_paragraph()
    p_disc = doc.add_paragraph()
    set_paragraph_spacing(p_disc, 12, 4)
    run_d = p_disc.add_run(
        'This consolidated assessment was prepared using the Marine Prosperity Index framework. '
        'The MPpI provides diagnostic guidance on investment priorities but does not prescribe '
        'specific intervention designs. Local stakeholder engagement, cultural sensitivity, and '
        'equity assessment should inform implementation of all recommended actions.'
    )
    run_d.italic = True
    run_d.font.size = Pt(9.5)
    run_d.font.color.rgb = COLOR_TEXT_GRAY

    doc.save(output_path)
    print(f'  Saved consolidated: {os.path.basename(output_path)}')


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('Building individual policy brief DOCX files...')
    for slug, (md_file, map_slug) in COMMUNITY_MAP.items():
        md_path = os.path.join(MD_DIR, md_file)
        if not os.path.exists(md_path):
            print(f'  SKIP (not found): {md_file}')
            continue
        out_path = os.path.join(OUT_DIR, f'policy_brief_{slug}.docx')
        # Extract community name from filename
        name = slug.replace('_', ' ').title()
        try:
            doc = build_docx(md_path, map_slug, name)
            doc.save(out_path)
            print(f'  Saved: policy_brief_{slug}.docx')
        except Exception as e:
            import traceback
            print(f'  ERROR {slug}: {e}')
            traceback.print_exc()

    print('\nBuilding consolidated regional report...')
    consolidated_path = os.path.join(OUT_DIR, 'consolidated_report_gulf_of_california.docx')
    try:
        build_consolidated(consolidated_path)
    except Exception as e:
        import traceback
        print(f'  ERROR consolidated: {e}')
        traceback.print_exc()

    print('\nAll done.')
